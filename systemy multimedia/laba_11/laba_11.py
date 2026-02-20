from io import BytesIO
import cv2
import numpy as np
import os
import matplotlib.pyplot as plt
from skimage.metrics import peak_signal_noise_ratio, structural_similarity
from docx import Document
from docx.shared import Inches

PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resources")

raport = Document()


def water_mark(img, mask, alpha=0.25):
    assert (img.shape[0] == mask.shape[0]) and (
        img.shape[1] == mask.shape[1]
    ), "Wrong size"
    if len(img.shape) < 3:
        flag = True
        t_img = cv2.cvtColor(img, cv2.COLOR_GRAY2RGBA)
    else:
        flag = False
        t_img = cv2.cvtColor(img, cv2.COLOR_RGB2RGBA)
    if mask.dtype == bool:
        t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_GRAY2RGBA)
    elif mask.dtype == np.uint8:
        if len(mask.shape) < 3:
            t_mask = cv2.cvtColor((mask).astype(np.uint8), cv2.COLOR_GRAY2RGBA)
        else:
            t_mask = cv2.cvtColor((mask).astype(np.uint8), cv2.COLOR_RGB2RGBA)
    else:
        if len(mask.shape) < 3:
            t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_GRAY2RGBA)
        else:
            t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_RGB2RGBA)
    t_out = cv2.addWeighted(t_img, 1, t_mask, alpha, 0)
    if flag:
        out = cv2.cvtColor(t_out, cv2.COLOR_RGBA2GRAY)
    else:
        out = cv2.cvtColor(t_out, cv2.COLOR_RGBA2RGB)
    return out


def put_data(img, data, binary_mask=np.uint8(1)):
    assert img.dtype == np.uint8, "img wrong data type"
    assert binary_mask.dtype == np.uint8, "binary_mask wrong data type"
    un_binary_mask = np.unpackbits(binary_mask)
    if data.dtype != bool:
        unpacked_data = np.unpackbits(data)
    else:
        unpacked_data = data
    dataspace = img.shape[0] * img.shape[1] * np.sum(un_binary_mask)
    assert dataspace >= unpacked_data.size, "too much data"
    if dataspace == unpacked_data.size:
        prepered_data = unpacked_data.reshape(
            img.shape[0], img.shape[1], int(np.sum(un_binary_mask))
        ).astype(np.uint8)
    else:
        prepered_data = np.resize(
            unpacked_data, (img.shape[0], img.shape[1], int(np.sum(un_binary_mask)))
        ).astype(np.uint8)
    mask = np.full((img.shape[0], img.shape[1]), binary_mask)
    img = np.bitwise_and(img, np.invert(mask))
    bv = 0
    for i, b in enumerate(un_binary_mask[::-1]):
        if b:
            temp = prepered_data[:, :, bv]
            temp = np.left_shift(temp, i)
            img = np.bitwise_or(img, temp)
            bv += 1
    return img


def pop_data(img, binary_mask=np.uint8(1), out_shape=None):
    un_binary_mask = np.unpackbits(binary_mask)
    data = np.zeros((img.shape[0], img.shape[1], np.sum(un_binary_mask))).astype(
        np.uint8
    )
    bv = 0
    for i, b in enumerate(un_binary_mask[::-1]):
        if b:
            mask = np.full((img.shape[0], img.shape[1]), 2**i)
            temp = np.bitwise_and(img, mask)
            data[:, :, bv] = temp[:, :].astype(np.uint8)
            bv += 1
    if out_shape != None:
        tmp = np.packbits(data.flatten())
        tmp = tmp[: np.prod(out_shape)]
        data = tmp.reshape(out_shape)
    return data


def load_image(img_path):
    img = cv2.imread(img_path)
    if img is None:
        raise FileNotFoundError(f"Image not found: {img_path}")
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

    return img


def load_text(text_path, max_len=None):
    with open(text_path, "r", encoding="utf-8") as f:
        text = f.read()

    if max_len is not None:
        text = text[:max_len]

    return text


def encode_text_to_image(img, text, binary_mask=np.uint8(1)):
    text_bytes = np.frombuffer(text.encode("utf-8"), dtype=np.uint8)
    img_encoded = img.copy()
    blue_channel = img_encoded[:, :, 2]
    blue_encoded = put_data(blue_channel, text_bytes, binary_mask)
    img_encoded[:, :, 2] = blue_encoded

    return img_encoded, text_bytes


def decode_text_from_image(img_encoded, binary_mask, text_bytes_shape):
    extracted_bits = pop_data(
        img_encoded[:, :, 2], binary_mask, out_shape=text_bytes_shape
    )

    extracted_bytes = extracted_bits.astype(np.uint8).tobytes()

    try:
        extracted_text = extracted_bytes.decode("utf-8")
    except UnicodeDecodeError:
        extracted_text = extracted_bytes.decode("utf-8", errors="replace")

    return extracted_text


def stego_encode(img_carrier, img_secret, masks):
    H, W, _ = img_carrier.shape
    encoded = img_carrier.copy()

    for layer, mask in enumerate(masks):
        mask = np.uint8(mask)

        positions = [i for i in range(8) if (mask >> i) & 1]
        secret_to_encode = np.zeros((H, W, len(positions)), dtype=bool)

        for idx, bit in enumerate(positions):
            secret_to_encode[:, :, idx] = (
                (img_secret[:, :, layer] << bit) & 128
            ).astype(bool)

        encoded[:, :, layer] = put_data(encoded[:, :, layer], secret_to_encode, mask)

    return encoded


def stego_decode(img_stego, masks):
    extracted = []
    for layer, mask in enumerate(masks):
        channel_data = pop_data(img_stego[:, :, layer], np.uint8(mask))
        channel_data = channel_data.sum(axis=2)

        positions = [i for i in range(8) if (mask << i) & 128]

        min_pos = min(positions)
        channel_data = np.left_shift(channel_data, min_pos)

        extracted.append(channel_data)

    return np.stack(extracted, axis=2).astype(np.uint8)


def compute_channel_metrics(img1, img2):
    psnr = [
        peak_signal_noise_ratio(img1[:, :, i], img2[:, :, i], data_range=255)
        for i in range(3)
    ]
    ssim = [
        structural_similarity(img1[:, :, i], img2[:, :, i], data_range=255)
        for i in range(3)
    ]
    return psnr, ssim


def compute_all_metrics(img1, img2):
    psnr = peak_signal_noise_ratio(img1, img2, data_range=255)
    ssim = structural_similarity(img1, img2, data_range=255, channel_axis=-1)
    return psnr, ssim


def visualize_stego_bitplanes(img_secret, masks):
    H, W, _ = img_secret.shape
    fig, axes = plt.subplots(4, 4, figsize=(12, 12))
    fig.suptitle("Deconstruction of Secret Image", fontsize=16)

    channel_names = list("RGB")
    bit_positions = [[i for i in range(8) if (masks[c] >> i) & 1] for c in range(3)]

    for c in range(3):
        pos = bit_positions[c]
        for idx, bit in enumerate(pos[:3]):
            weight = 128 >> bit  # 10000000 >> 0, 10000000 >> 1 = 01000000
            plane = ((img_secret[:, :, c] & weight) > 0).astype(np.uint8)
            ax = axes[c, idx]
            ax.imshow(plane, cmap="gray")
            ax.set_title(f"bits → {channel_names[c]} & {weight}")
            ax.axis("off")
        for idx in range(len(pos), 3):
            axes[c, idx].axis("off")

        reconstructed = np.zeros((H, W), dtype=np.uint8)
        for bit in pos:
            weight = 128 >> bit
            reconstructed += ((img_secret[:, :, c] & weight) > 0).astype(
                np.uint8
            ) * weight
        ax = axes[c, 3]
        color_layer = np.zeros((H, W, 3), dtype=np.uint8)
        color_layer[:, :, c] = reconstructed
        ax.imshow(color_layer)
        ax.set_title(f"reconstructed {channel_names[c]} layer")
        ax.axis("off")

    for i, c in enumerate([2, 0, 1]):
        pos = bit_positions[c]
        reconstructed = np.zeros((H, W), dtype=np.uint8)
        for bit in pos:
            reconstructed += ((img_secret[:, :, c] & (128 >> bit)) > 0).astype(
                np.uint8
            ) * (128 >> bit)
        ax = axes[3, i]
        ax.imshow(reconstructed, cmap="gray")
        ax.set_title(f"reconstructed {channel_names[c]} layer")
        ax.axis("off")

    final_rgb = np.zeros((H, W, 3), dtype=np.uint8)
    for c in range(3):
        for bit in bit_positions[c]:
            final_rgb[:, :, c] += ((img_secret[:, :, c] & (128 >> bit)) > 0).astype(
                np.uint8
            ) * (128 >> bit)
    ax = axes[3, 3]
    ax.imshow(final_rgb)
    ax.set_title("reconstructed Red+Green+Blue layer")
    ax.axis("off")

    plt.tight_layout(rect=[0, 0, 1, 0.96])
    plt.show()

def hide_image_in_rgb(carrier, secret, masks):
    encoded = carrier.copy()
    for ch, bits in zip([0, 1, 2], masks):
        carrier_ch = carrier[:, :, ch].copy()
        secret_ch = secret[:, :, ch]
        for b in range(bits):
            bit_mask = (secret_ch >> (7 - b)) & 1
            clear_mask = 255 - (1 << b)
            carrier_ch = (carrier_ch & clear_mask) | (bit_mask << b)
        encoded[:, :, ch] = carrier_ch
    return encoded

def recover_hidden_image(encoded, masks):
    recovered = np.zeros_like(encoded)
    for ch, bits in zip([0, 1, 2], masks):
        enc_ch = encoded[:, :, ch]
        for b in range(bits):
            recovered[:, :, ch] |= ((enc_ch >> b) & 1) << (7 - b)
    return recovered

def plot_comparison(original, stego, extracted, title_prefix=""):
    fig, axs = plt.subplots(2, 3, figsize=(12, 8))
    axs[0, 0].imshow(original)
    axs[0, 0].set_title(f"Original {title_prefix}")
    axs[0, 1].imshow(stego)
    axs[0, 1].set_title(f"Stego {title_prefix}")
    axs[0, 2].imshow(extracted)
    axs[0, 2].set_title(f"Extracted secret {title_prefix}")
    for i, name in enumerate(["R", "G", "B"]):
        axs[1, i].imshow(
            np.abs(original[:, :, i].astype(int) - stego[:, :, i].astype(int)),
            cmap="hot",
        )
        axs[1, i].set_title(f"Diff {name}")
    for ax in axs.flat:
        ax.axis("off")
    plt.tight_layout()

    return fig


def compare_top_bits(arr1: np.ndarray, arr2: np.ndarray, n: int) -> np.ndarray:
    assert arr1.shape == arr2.shape, "Shapes must match"
    assert (
        arr1.dtype == np.uint8 and arr2.dtype == np.uint8
    ), "Arrays must be dtype uint8"
    assert 1 <= n <= 8, "n must be in [1,8]"

    shift = 8 - n
    top1 = np.right_shift(arr1, shift)
    top2 = np.right_shift(arr2, shift)
    return top1 == top2


def task2():
    image_name = "image1.png"
    text_file_name = "text.txt"
    binary_mask = np.uint8(1)

    img = load_image(os.path.join(PATH, image_name))
    text = load_text(os.path.join(PATH, text_file_name))

    img_encoded, text_bytes = encode_text_to_image(img, text, binary_mask)
    extracted_text = decode_text_from_image(img_encoded, binary_mask, text_bytes.shape)

    psnr, ssim = compute_channel_metrics(img, img_encoded)

    raport.add_heading("Ukrywanie tekstu w niebieskim kanale (LSB)", level=2)
    raport.add_paragraph(f"Tekst oryginalny: {text}")
    raport.add_paragraph(f"Tekst odzyskany: {extracted_text}")
    equal = "Tak" if text == extracted_text else "Nie"
    raport.add_paragraph(f"Czy teksty są identyczne: {equal}")
    table = raport.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Kanał"
    hdr_cells[1].text = "PSNR"
    hdr_cells[2].text = "SSIM"
    row_cells = table.rows[1].cells
    row_cells[0].text = "B"
    row_cells[1].text = f"{psnr[2]:.2f}"
    row_cells[2].text = f"{ssim[2]:.4f}"
    raport.add_page_break()


def task3():
    carrier_name = "image1.png"
    secret_name = "image2.png"
    masks = [3, 3, 7]

    img_carrier = load_image(os.path.join(PATH, carrier_name))
    img_secret = load_image(os.path.join(PATH, secret_name))

    # encoded = stego_encode(img_carrier, img_secret, masks)
    # decoded = stego_decode(encoded, masks)
    encoded = hide_image_in_rgb(img_carrier, img_secret, masks)
    decoded = recover_hidden_image(encoded, masks)
    
    extracted_img = np.clip(decoded, 0, 255).astype(np.uint8)
    psnr, ssim = compute_channel_metrics(img_carrier, encoded)
    fig = plot_comparison(img_carrier, encoded, extracted_img)
    
    raport.add_heading(f"Ukrywanie kolorowego obrazu {masks=}", level=2)
    table = raport.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Kanał"
    hdr_cells[1].text = "PSNR"
    hdr_cells[2].text = "SSIM"

    for i, c in enumerate(["R", "G", "B"]):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = c
        row_cells[1].text = f"{psnr[i]:.2f}"
        row_cells[2].text = f"{ssim[i]:.4f}"

    memfile = BytesIO()
    fig.savefig(memfile)
    plt.close(fig)
    raport.add_paragraph()
    raport.add_picture(memfile, width=Inches(6))
    raport.add_page_break()


def task4():
    carrier_name = "image1.png"
    secret_name = "image2.png"

    img_carrier = load_image(os.path.join(PATH, carrier_name))
    img_secret = load_image(os.path.join(PATH, secret_name))

    # bits_range = range(1, 9)
    # all_masks = [[(1 << b) - 1] * 3 for b in bits_range]
    bits_range = range(1, 8)
    all_masks = [[b] * 3 for b in bits_range]
    

    results = []
    figs = []

    for masks in all_masks:
        # encoded = stego_encode(img_carrier, img_secret, masks)
        # decoded = stego_decode(encoded, masks)
        encoded = hide_image_in_rgb(img_carrier, img_secret, masks)
        decoded = recover_hidden_image(encoded, masks)

        extracted_img = np.clip(decoded, 0, 255).astype(np.uint8)
        psnr, ssim = compute_channel_metrics(img_carrier, encoded)
        results.append((masks, psnr, ssim))
        fig = plot_comparison(
            img_carrier, encoded, extracted_img, title_prefix=f"{masks=}"
        )
        figs.append(fig)

    raport.add_heading("Badanie budżetu bitowego", level=2)
    table = raport.add_table(rows=len(results) + 1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Maska"
    hdr[1].text = "PSNR R"
    hdr[2].text = "PSNR G"
    hdr[3].text = "PSNR B"
    hdr[4].text = "SSIM R"
    hdr[5].text = "SSIM G"
    hdr[6].text = "SSIM B"
    for i, (masks, psnr_vals, ssim_vals) in enumerate(results):
        row = table.rows[i + 1].cells
        row[0].text = str(masks)
        row[1].text = f"{psnr_vals[0]:.2f}"
        row[2].text = f"{psnr_vals[1]:.2f}"
        row[3].text = f"{psnr_vals[2]:.2f}"
        row[4].text = f"{ssim_vals[0]:.4f}"
        row[5].text = f"{ssim_vals[1]:.4f}"
        row[6].text = f"{ssim_vals[2]:.4f}"

    psnr_table = np.array([results[i][1] for i in range(len(results))])
    ssim_table = np.array([results[i][2] for i in range(len(results))])
    plt.figure(figsize=(10, 4))
    plt.subplot(1, 2, 1)
    for i, c in enumerate("RGB"):
        plt.plot(bits_range, psnr_table[:, i], label=f"PSNR {c}")
    plt.xlabel("Bits per channel")
    plt.ylabel("PSNR")
    plt.legend()
    plt.subplot(1, 2, 2)
    for i, c in enumerate("RGB"):
        plt.plot(bits_range, ssim_table[:, i], label=f"SSIM {c}")
    plt.xlabel("Bits per channel")
    plt.ylabel("SSIM")
    plt.legend()
    plt.tight_layout()

    memfile = BytesIO()
    plt.savefig(memfile)
    plt.close()
    raport.add_paragraph()
    raport.add_picture(memfile, width=Inches(6))

    for fig in figs:
        memfile = BytesIO()
        fig.savefig(memfile)
        plt.close(fig)
        raport.add_paragraph()
        raport.add_picture(memfile, width=Inches(6))

    raport.add_page_break()


def task5():
    carrier_name = "image1.png"
    watermark_name = "logo.png"
    alphas = [0.1, 0.25, 0.5]

    img_carrier = load_image(os.path.join(PATH, carrier_name))
    watermark_mask = load_image(os.path.join(PATH, watermark_name))

    if watermark_mask.ndim == 3:
        watermark_mask = cv2.cvtColor(watermark_mask, cv2.COLOR_RGB2GRAY)

    watermark_mask = watermark_mask > 127

    raport.add_heading("Znak wodny", level=2)
    table = raport.add_table(rows=len(alphas) + 1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Alpha"
    hdr_cells[1].text = "PSNR"
    hdr_cells[2].text = "SSIM"

    for i, alpha in enumerate(alphas):
        watermarked = water_mark(img_carrier.copy(), watermark_mask, alpha=alpha)
        psnr, ssim = compute_all_metrics(img_carrier, watermarked)

        plt.figure(figsize=(10, 4))
        plt.subplot(1, 2, 1)
        plt.imshow(img_carrier)
        plt.title("Original Carrier")
        plt.axis("off")
        plt.subplot(1, 2, 2)
        plt.imshow(watermarked)
        plt.title(f"Watermarked Image ({alpha=})")
        plt.axis("off")
        plt.tight_layout()

        row_cells = table.rows[i + 1].cells
        row_cells[0].text = f"{alpha:.2f}"
        row_cells[1].text = f"{psnr:.2f}"
        row_cells[2].text = f"{ssim:.4f}"

        memfile = BytesIO()
        plt.savefig(memfile)
        plt.close()
        raport.add_paragraph()
        raport.add_picture(memfile, width=Inches(6))

    raport.add_page_break()


if __name__ == "__main__":
    task2()
    task3()
    task4()
    task5()
    raport.save("raport.docx")

